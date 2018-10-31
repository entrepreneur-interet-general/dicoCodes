# -*- coding: utf-8 -*-
"""
Created on Fri Mar 16 18:13:54 2018

@author: samah.ghalloussi
"""


import sys, copy
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
import docutils
import logging
logging.basicConfig(level=logging.DEBUG)

"""
Module that extract text from MS XML Word document (.docx).
(Inspired by python-docx <https://github.com/mikemaccana/python-docx>)
"""


def rmq_process(rmq):
    """ Return structured data from remarks content in CARE-S : 
        rawtext and dict of choices
    """
#    separators = [':', '.', '\t', '=']
    rmq = rmq.replace('\u2022','')
    separators = {':':0, '.':0, '\t':0, '=':0, '\n':0}
    dico = {}
    for sep in separators :
        separators[sep] = len(rmq.split(sep)) - 1
        if separators[sep] > 0 and sep != '\n' :
            tempodico = {}
            tempotext = ''
            for line in rmq.split('\n') :
                val = ''
                # handle most of the cases
                if len(line.split(sep)) == 2 :
                    linecontent = line.split(sep)
                    val = linecontent[0].strip()
                    desc = linecontent[1].strip()
                    # handle =0 si
                    if val == '':
                        linecontent = desc.split(' ')
                        val = linecontent[0]
                        desc = ' '.join(linecontent[1::])
                # handle ... with sep='.'
                elif len(line.split(sep)) > 2 :
                    val = line.split(sep)[0].strip()
                    linecontent = line.split(' ')
                    desc = ' '.join(linecontent[1::])
                # handle space separator
                elif line.split(' ')[0].isdigit() :
                    linecontent = line.split(' ')
                    val = linecontent[0]
                    desc = ' '.join(linecontent[1::])
                  
                repval = val.replace('-',' ')
                #strip to handle 97-10
                if repval.strip().isdigit() :
                    tempodico[int(val)] = desc.replace('?', '').strip()
                else :
                    tempotext += line + '\n'
            if tempodico != {} :
                dico[sep] = tempodico
                dico[sep]['text'] = '' ###tempotext
    
    #2nd process
    textRemarques = ''
    newdico = {}
    for sepdico in dico :
        for e in dico[sepdico] :
            if e == 'text' :
                textRemarques = dico[sepdico]['text']
            else :
                newdico[int(e)] = dico[sepdico][e]                           
                    
    return textRemarques, newdico

def check_lines(dicotable):
    """ Check count of columns in tables from CARE-S specs
    """
    count = {}
    for line in dicotable :
        try : count[len(dicotable[line])] += 1
        except: count[len(dicotable[line])] = 1
    
    if len(count) == 1:
        return True
    else :
        return False
            
def parse_spec(tagslist):
    """ Return a tempo dictionary with extracted data from input document tagslist
    """
    current = {'tbl':0}
    i, r, c = 0, 0, 0
    table = {}
    tmp = ''
    title, desc = [], ''
    intable, soustable = False, False
    istitle = False
    for line in tagslist:
        if line.startswith('<p ') :
            pass
        elif line == '</p>' :
            if intable :
                tmp += '\n'
            else :
                desc += '\n'
            istitle = False
        elif line == '<pStyle val="Titre1" />':
            istitle = True
        elif line == '<pStyle val="Titre2" />':
            istitle = True
        elif line == '<pStyle val="Titre3" />':
            istitle = True
        elif line == '<tbl>' :
            current['tbl'] += 1
            if not intable :
                intable = True
                r = 0
                table[i] = {'table':{}, 'atitle':title, 'desc':desc.strip()}
                title, desc = [], ''
            else :
                soustable = True
            
        elif line.startswith('<tr ') and not soustable :
            table[i]['table'][r] = {}
            c = 0
            
        elif line == '<tc>' and not soustable :
            tmp = ''
            
        elif line == '</tbl>' :
            current['tbl'] += -1
            if not soustable :
                intable = False
                i += 1
            else :
                soustable = False
            
        elif line == '</tr>' and not soustable :
            r += 1
        
        elif line == '</tc>' and not soustable :
#            table[i]['table'][r][c] = tmp[0:-1]
            table[i]['table'][r][c] = tmp.strip().replace('\n(nom de la balise)', '').replace(' en sortie de l\u2019identification','')
            tmp = ''
            c += 1
            
        elif line == '<tab />' :
            tmp += '\t'
            
        elif line.endswith('</t>') :
            text = line.replace('</t>', '').replace('\u00a0', ' ')
            if istitle :
#                if text != 'Tables' :
                    title.append(text.strip())
                
            elif intable :
                tmp += text
            else :
                desc += text
    
    logging.debug("parse_spec OK")
    
    return table

def analyse_spec(table):
    """ Return a dictionary of parsed data from dict table of extracted data
    """
    newtable = {}
    for num in table :
      if check_lines(table[num]['table']):
        try : atitle = table[num]['atitle'][-1]
        except : atitle = ''
        newtable[atitle] = {}
        for irow, row in enumerate(table[num]['table']) :
            if irow > 0 :
                addico = {'row':row, 'itable': num, 'Table':table[num]['atitle']}
                for col in table[num]['table'][row] :
                    addico[table[num]['table'][0][col]] = table[num]['table'][row][col]
                    
                try :
                    textRemarques, newdico = rmq_process(addico['Remarques'])
                    addico['RemarquesTexte'] = textRemarques
                    addico['RemarquesValeurs'] = newdico
                except Exception as e:
                    logging.debug("error rmq", str(e))
                    
                try : 
                    vartab = addico['Nom variable'].split()
                    #gerer 
                    if '\u00e0' in vartab and len(vartab)==3:
                        #gerer HANDIC2E_A à HANDIC2E_T
                        #if to do
                        if '_C' in vartab[0] :
                            target = '_C'
                            addplus = 2
                        elif '_' in vartab[0] :
                            target = '_'
                            addplus = 1
                        else :
                            target = '_'
                            addplus = 0
                        varpos = vartab[0].rfind(target) + addplus 
                        var = vartab[0][0:varpos]
                        try : startnum = int(vartab[0][varpos::])
                        except : startnum = 1 #temposolution
                        endnumpos = vartab[2].rfind(target) + addplus
                        try : endnum = int(vartab[2][endnumpos::])
                        except : endnum = 2 #temposolution
                        ext = ''
                        if endnum > 9 :
                            ext = '0'
                        for i in range(endnum):
                            number = i+startnum
                            if ext == '0' and number > 9 :
                                ext = ''
                            varname = var + ext + str(number)
#                            newtable[num].append({varname:addico})
                            newtable[atitle][varname] = copy.deepcopy(addico)
                            newtable[atitle][varname]['duplicate'] = True
                            newtable[atitle][varname]['duplinb'] = number
                            
                    # handle 'NATIO1N1 et NATIO1N2'
                    elif ' et ' in vartab and len(vartab)==3:
                        newtable[atitle][vartab[0]] = copy.deepcopy(addico)
                        newtable[atitle][vartab[2]] = copy.deepcopy(addico)
                        newtable[atitle][vartab[0]]['duplicate'] = True
                        newtable[atitle][vartab[2]]['duplicate'] = True
                        newtable[atitle][vartab[0]]['duplinb'] = 1
                        newtable[atitle][vartab[2]]['duplinb'] = 2
                        
                    else :
                        newtable[atitle][vartab[0]] = addico
                       
                except Exception as e: 
                    logging.debug("error vartab:", str(e))
    
    logging.debug("analyse_spec OK")
    
    return newtable 
    
def list_variables(tables):
    """ Return a list of all the var names extracted from tables dictionary
    """
    variables = []
    for nametable in tables :
        for variable in tables[nametable] :
            variables.append(variable)
    
    logging.debug("list_variables OK")
    return variables
    
def parse_quest(tagslist, variables):
    """ Return structured data from CARE-S questionnary input
    """
    dicoquest = {}
    paratext = ''
    variable = ''
    firstlinetext = False
    keepquest = False
    isvar = False
    iscadre = False 
#    isenum, istitle = False, False
    isblue, wasblue = False, False
    isbold, wasbold = False, False
#    dicoenum = {}
#    enumnb = 0
    for line in tagslist:
        if line.startswith('<p ') :
            firstlinetext = True
#            isenum, istitle = False, False
            wasbold, wasblue = False, False
            isbold = False
            isblue = False
            paratext = ''
        elif line.startswith('<r ') :
            if isblue :
                isblue = False
                wasblue = True
            else :
                wasblue = False
                
            if isbold :
                isbold = False
                wasbold = True
            else :
                wasbold = False
        elif line.startswith('</r>') :
            if isblue :
                isblue = False
                wasblue = True
            if isbold :
                isbold = False
                wasbold = True
        elif line == '<b />' :
            isbold = True

        elif line.startswith('<color val="0000FF" />')  :
            isblue = True
#        elif line == '<pStyle val="Titre2" />':
#            istitle = True
#        elif line.startswith('<ilvl val=') and not istitle and not iscadre and not line.endswith('"12" />') and not line.endswith('"0" />'):
#            isenum = True
        elif line == '<txbxContent>' :
            iscadre = True
        elif line == '</txbxContent>' :
            iscadre = False            

        elif line == '</p>' :
            if iscadre :
                pass
            elif paratext.strip() == '' :
                isvar = False
            elif paratext.strip() in variables or paratext.strip()+'_01' in variables or paratext.strip()+'_1' in variables :
                if variable != '' :
                    textRemarques, dicoquest[variable]['modalites'] = rmq_process(dicoquest[variable]['other'])
                    if textRemarques != '' :
                        dicoquest[variable]['textRemarques'] = textRemarques
#                    if dicoenum != {}:
#                        dicoquest[variable]['dicoenum'] = dicoenum
#                        dicoenum = {}
#                        enumnb = 0
#                        isenum = False
                variable = paratext.strip()
                dicoquest[variable] = {'other':'', 'instructions':''}
                keepquest = True
                isvar = True
                paratext = ''
            elif keepquest :
#                if isbold or wasbold :
                    dicoquest[variable]['Question'] = paratext.strip()
                    keepquest = False
#                else :
#                    dicoquest[variable]['infos'] = paratext + '\n'
            elif isvar :
                if (wasblue or isblue or isbold or paratext.startswith('Exemples')\
                        or paratext.startswith('Instruction') \
                        or paratext.startswith('(plusieurs') \
                        or "NSP/RF non autorisés" in paratext \
                        or paratext.startswith(' carte ')) \
                        and "(ne pas lire" not in paratext:
                    dicoquest[variable]['instructions'] += paratext + '\n'
#                    print(variable, isbold, paratext)
                else :
                    dicoquest[variable]['other'] += paratext + '\n'
        
        elif line.endswith('</t>') :
            if not iscadre :
                text = line.replace('</t>', '').replace('\u2026', ' ').replace('\u00a0', ' ')
                paratext += text
#                if isenum :
#                    if firstlinetext :
#                        enumnb += 1
#                        dicoenum[enumnb] = [text]
#                    else :
#                        dicoenum[enumnb][-1] += text
                firstlinetext = False
            
    logging.debug("parse_quest OK")
    
    return dicoquest


def merge_data(tables, addinfos):
    """ Return final tables with additionnal information from the questionnary input
    """
    i, j = 0, 0
    for nametable in tables :
        for variable in tables[nametable] :
            try :
#                if tables[nametable][variable]["duplicate"] == True :
                choice = tables[nametable][variable]["duplinb"]
                var = variable[0:variable.find('_')]
                try : 
                    add = addinfos[var]["modalites"][str(choice)]
                    iend = add.find('\uf0e8')
                    if iend != -1 :
                        tables[nametable][variable]["add"] = add[0:iend].strip()
                    else :
                        tables[nametable][variable]["add"] = add.strip()
                    i+=1
                except:
                    pass
            except :
                try : 
                    tables[nametable][variable]["addQuestion"] = addinfos[variable]["Question"]
                    tables[nametable][variable]["addModalites"] = addinfos[variable]["modalites"]
                    j+=1
                except :
                    pass
    
    return tables

def process_quest(inpath, variables):
    """ Return a dictionary with structured data from the questionnary input
    """
    tagslist = docutils.extract_tags(inpath)
    dicoquest = parse_quest(tagslist, variables)
    
    return dicoquest

def process_spec(inpath):
    """ Return a dictionary with structured data from the specs document input
    """
    tagslist = docutils.extract_tags(inpath)
    table = parse_spec(tagslist)
    dicospec = analyse_spec(table)
    
    return dicospec  

def process_data(inputspec, inputquest):
    """ Return a dictionary with structured data from the specs and questionnary inputs
    """
    dicospec = process_spec(inputspec)
    variables = list_variables(dicospec)
    dicoquest = process_quest(inputquest, variables)
    
    return merge_data(dicospec, dicoquest)
    


def set_styles(document):
    """ Define specific styles for CARE-S document output
    TO DO: format YAML
    """  
    styleHeading0 = document.styles['Title']
    styleHeading0.font.name = 'Calibri'
    styleHeading0.font.color.theme_color = None
    styleHeading0.font.size = Pt(14)
    styleHeading0.font.bold = True
    styleHeading0.paragraph_format.space_after = Pt(0)
    styleHeading0.paragraph_format.widow_control = True
    
    styleHeading1 = document.styles['Heading 1']
    styleHeading1.font.name = 'Cambria (Titres)'

    styleHeading2 = document.styles['Heading 2']
    styleHeading2.font.name = 'Cambria (Titres)'
    styleHeading2.font.size = Pt(13)
    styleHeading2.paragraph_format.left_indent = Inches(0.5)
    
    styleHeading3 = document.styles['Heading 3']
    styleHeading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    styleTableTitle = document.styles.add_style('TableTitle', WD_STYLE_TYPE.PARAGRAPH)
    styleTableTitle.font.name = 'Calibri'
    styleTableTitle.font.size = Pt(20)
    styleTableTitle.font.italic = False
    styleTableTitle.font.bold = True
    styleTableTitle.paragraph_format.space_after = Pt(0)
    styleTableTitle.paragraph_format.widow_control = True
    styleTableTitle.paragraph_format.keep_with_next = True
    
    styleVariable = document.styles.add_style('VariableCell', WD_STYLE_TYPE.PARAGRAPH)
    styleVariable.font.name = 'Calibri'
    styleVariable.font.size = Pt(14)
    styleVariable.font.italic = False
    styleVariable.font.bold = True
    styleVariable.paragraph_format.space_after = Pt(0)
    styleVariable.paragraph_format.widow_control = True
    styleVariable.paragraph_format.keep_with_next = True
    
    styleTableCell = document.styles.add_style('TableCell', WD_STYLE_TYPE.PARAGRAPH)
    styleTableCell.font.name = 'Calibri'
    styleTableCell.font.size = Pt(11)
    styleTableCell.paragraph_format.space_after = Pt(0)
    styleTableCell.paragraph_format.widow_control = True
    styleTableCell.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    styleTableCell.paragraph_format.keep_with_next = True
    
    styleLibelle = document.styles.add_style('LibelleCell', WD_STYLE_TYPE.PARAGRAPH)
    styleLibelle.font.name = 'Calibri'
    styleLibelle.font.size = Pt(12)
    styleLibelle.font.italic = False
    styleLibelle.font.bold = True
    styleLibelle.paragraph_format.space_before = Pt(0)
    styleLibelle.paragraph_format.space_after = Pt(0)
    styleLibelle.paragraph_format.widow_control = True
    styleLibelle.paragraph_format.keep_with_next = True
    
    stylePrimaireCell = document.styles.add_style('PrimaireCell', WD_STYLE_TYPE.PARAGRAPH)
    stylePrimaireCell.font.name = 'Calibri'
    stylePrimaireCell.font.size = Pt(11)
    stylePrimaireCell.paragraph_format.space_before = Pt(6)
    stylePrimaireCell.paragraph_format.space_after = Pt(6)
    stylePrimaireCell.paragraph_format.widow_control = True
    stylePrimaireCell.paragraph_format.keep_with_next = True
    
    styleNormalCell = document.styles.add_style('NormalCell', WD_STYLE_TYPE.PARAGRAPH)
    styleNormalCell.font.name = 'Calibri'
    styleNormalCell.font.size = Pt(11)
    styleNormalCell.font.italic = False
    styleNormalCell.font.bold = False
    styleNormalCell.paragraph_format.space_before = Pt(0)
    styleNormalCell.paragraph_format.space_after = Pt(0)
    styleNormalCell.paragraph_format.widow_control = True
    styleNormalCell.paragraph_format.keep_with_next = True
    
    return document


def add_headings(document) :
    """ Add specific table of content titles in the document output
    """
    document.add_heading('I. Présentation de l’enquête', 1)
    document.add_heading('II. Échantillonnage et pondérations', 1)
    document.add_heading('III. Liste des variables', 1)    
    document.add_heading('IV. Présentation des traitements aval', 1)
#    document.add_heading('A. Imputation des variables', 2)
#    document.add_heading('B. Score de santé mentale : variable MH', 2)
#    document.add_heading('C. Deux estimations possibles du GIR : variables groupelarge et grouperestreint', 2)
    document.add_heading('V. Dictionnaire des codes', 1)
    document.add_page_break()
    
    return document
    
def add_firstpage(document) :
    """ Add the firt page content in the document output
    """
    p = document.add_paragraph('\n\n', style='ArialTitle')
    p.add_run('C', 'ArialBlack')
    p.add_run('apacités, ')
    p.add_run('A', 'ArialBlack')
    p.add_run('ides et ')
    p.add_run('RE', 'ArialBlack')
    p.add_run('ssources\ndes seniors\n\n\n')

    p1 = document.add_paragraph('Volet « seniors »', style='ArialTitle')
    p1.add_run('\n\nDictionnaire des codes\n\n')
    document.add_paragraph('Version du ' + docutils.get_date() + ' : générée automatiquement', style='Desc')
    p2 = document.add_paragraph('\n\n\n', style='ArialTitle')
    p2.add_run('\nEnquête 2015', 'ArialDesc')
    document.add_page_break()
    
    return document



 

def add_dico(document, tables):
    """ Write CARE-S data into the new document 
    """
    loggerTable = logging.getLogger("Table")
    # cell border values by default
    cellbr = {"sz": 5, "val": "single"}
    # list of table names to rename as SENIOR table :
    groupedtables = ["FAMILLE", "TUTELLES", "SANTE", "LIMITATIONS", "AIDESTECH", "SOINSHEBERGEMENT", "RESTRICTIONS", "ACCESSIBILITE", "DEPENSES", "AIDEHUMAINE"]
    for name in tables :
        loggerTable.debug(name)
        if name in groupedtables :
            paragraph = document.add_paragraph(name, style='TableTitle')
            tableName = 'Table : SENIORS'
        else :
            tableName = 'Table : '+ name
            paragraph = document.add_paragraph(tableName, style='TableTitle')
            
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for variable in tables[name] :
            table = document.add_table(rows=0, cols=3)
            # set column widths
            table.columns[0].width = Inches(1.0)
            table.columns[1].width = Inches(5.0)
            table.columns[2].width = Inches(2.0)

            cells = table.add_row().cells
            cells[0].add_paragraph()
            cells[1].merge(cells[2])
            cells[0].merge(cells[1])
            
            cells = table.add_row().cells
            docutils.cleanpara(cells[0], linespacing=WD_LINE_SPACING.SINGLE).add_paragraph(variable, style='VariableCell')
            
            docutils.cleanpara(cells[2]).add_paragraph(tableName, style='TableCell')
            cells[0].merge(cells[1])
            
            rmqdefault = tables[name][variable]["Remarques"].strip()
            try : rmqdico = tables[name][variable]["addModalites"]
            except : rmqdico = tables[name][variable]["RemarquesValeurs"]
            
            try : libelle = tables[name][variable]["addQuestion"].strip()
            except: libelle = tables[name][variable]["Libell\u00e9"].strip()
            
            try :
                libelle += '\n... ' + tables[name][variable]["add"]
            except :
                pass
            if libelle != '' :
                cells = table.add_row().cells
                docutils.cleanpara(cells[0]).add_paragraph(libelle, style='LibelleCell')
                cells[1].merge(cells[2])
                cells[0].merge(cells[1])

            if rmqdico == {}:
                if "Codage au format 1/0" in rmqdefault :
                    rmqdefault = rmqdefault.replace("Codage au format 1/0", "").strip()
                    rmqdico = {1:'Oui', 0:'Non'}
                    
            try : 
                nb = tables[name][variable]["Nb de positions"]
                if nb != '1' and nb != '' and rmqdico == {} :
                    cells = table.add_row().cells
                    cells[0].merge(cells[1])
                    cells = table.add_row().cells
                    cells[0].text = 'de 1 à ' + tables[name][variable]["Nb de positions"]
                    docutils.set_cell_border(cells[0], top=cellbr, bottom=cellbr, start=cellbr, end=cellbr)
                    
            except :
                pass
            
            if rmqdico != {} :
                cells = table.add_row().cells
                cells[0].merge(cells[1])
                for val in rmqdico :
                    cells = table.add_row().cells
                    docutils.set_cell_border(cells[0], top=cellbr, bottom=cellbr, start=cellbr, end=cellbr)
                    cells[0].text = str(val)
                    docutils.set_cell_border(cells[1], top=cellbr, bottom=cellbr, start=cellbr, end=cellbr)
                    cells[1].text = rmqdico[val]
            try : 
                valeurs = tables[name][variable]["Valeurs"]
                cells = table.add_row().cells
                docutils.cleanpara(cells[0]).add_paragraph(valeurs, style='PrimaireCell')
                docutils.cleanpara(cells[1]).add_paragraph("PRIMAIRE", style='PrimaireCell')
            except :
                pass
            
            rmqtexte = tables[name][variable]["RemarquesTexte"].strip()
            if rmqdico == {} and rmqtexte == '' and rmqdefault != '' :
                cells = table.add_row().cells
                cells[1].merge(cells[2])
                cells[0].merge(cells[1])
                cells[0].text = rmqdefault
            elif rmqtexte != '' :
                cells = table.add_row().cells
                cells[1].merge(cells[2])
                cells[0].merge(cells[1])
                cells[0].text = rmqtexte
                
        document.add_page_break()
        
    return document


def create_doc(tables, outputdoc):
    """ Create a new document for CARE-S
        For more information, see: https://python-docx.readthedocs.io/en/latest/
    """
    document = docutils.init_document()
    document = set_styles(document)
    document = add_firstpage(document)
    document = add_headings(document)
    
    document = add_dico(document, tables)
    
    document.add_heading('VI. Index', 1)
    document.save(outputdoc)
    
    logging.debug('create_doc OK')


if __name__ == '__main__':
    # To run the program, type-in $ python dicoCodes_CARES.py [inputspec] [inputquest] [outpath]
    try :
        inputspec = sys.argv[1]
        inputquest = sys.argv[2]
        outpath = sys.argv[3]
        details = False #sys.argv[4]
    except :
        inputspec = 'tests/inputs/Specifications_CARES.docx'
        inputquest = 'tests/inputs/Questionnaire_CARES.docx'
        outpath = 'tests/outputs/'
        details = True

    if not outpath.endswith('/') :
        outpath += '/'
        
    dicodata = process_data(inputspec, inputquest)

    if details :
        docutils.docxtoxml(inputquest, outpath +inputquest.split('/')[-1]+"-xmlcontent.xml")
        docutils.write_json(dicodata, outpath + 'dicodata-CARES.json')

    outputdoc = outpath + 'autoDictionnaireCARES-'+ str(docutils.date.today()) + '.docx'
    create_doc(dicodata, outputdoc)



