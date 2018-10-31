# -*- coding: utf-8 -*-
"""
Created on Mon Oct 29 12:16:55 2018

@author: samah.ghalloussi
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date
import json, traceback, zipfile
import logging
logging.basicConfig(level=logging.DEBUG)

try:
    from xml.etree.cElementTree import XML, tostring, tostringlist
except ImportError:
    from xml.etree.ElementTree import XML, tostring, tostringlist


def get_date():
    """ Return today's date with the month in letters
    """
    mois = ["Janvier", u"Février", "Mars", "Avril", "Mai", "Juin", "Juillet", u"Août", "Septembre", "Octobre", "Novembre", u"Décembre"]
    today = date.today()
    return str(today.day) + ' '+ mois[today.month - 1] + ' ' + str(today.year)


def cleanpara(cell, linespacing = None):
    """ Return a clear cell by removing the paragraph added automatically
        when a row is added to a table
    """
    paragraph = cell.paragraphs[-1]
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
    if linespacing != None :
        paragraph.line_spacing_rule = linespacing
        
    return cell

def add_italic(target, text, start_trigger='[', end_trigger=']', style='normalCell'):
    """ Add italic to parts of text defined between a specific trigger
    """
    para = target.add_paragraph(style=style)
    #add italic only if there is the same number of end and start triggers
    if text.count(start_trigger) > 0 and text.count(start_trigger) == text.count(start_trigger) :
        for section in text.split(start_trigger) :
            if section.endswith(end_trigger) :
                run = para.add_run(start_trigger + section)
                run.font.italic = True
            elif end_trigger in section :
                run = para.add_run(start_trigger + section.split(end_trigger)[0] + end_trigger)
                run.font.italic = True
                run = para.add_run(section.split(end_trigger)[1])
                run.font.italic = False
            else :
                run = para.add_run(section)
                run.font.italic = False
    else :
        para.add_run(text)
    
    return para

def add_color(rgb, start_trigger, end_trigger, target, text, style='normalCell'):
    """ Change color of the text defined between a specific trigger
    """
    para = target.add_paragraph(style=style)
    for section in text.split(start_trigger) :
        if section.endswith(end_trigger) :
            run = para.add_run(section.replace(end_trigger, ''))
            run.font.color.rgb = rgb
        elif end_trigger in section :
            run = para.add_run(section.split(end_trigger)[0])
            run.font.color.rgb = rgb
            run = para.add_run(section.split(end_trigger)[1])
            run.font.color.rgb = RGBColor(0, 0, 0)
        else :
            run = para.add_run(section)
            run.font.color.rgb = RGBColor(0, 0, 0)

def init_document():
    """ Create the document output
    """
    document = Document()
    core_properties = document.core_properties
    core_properties.author = 'DREES'
    core_properties.language = 'fr'
    
    document = set_primary_styles(document)
    
    logging.debug('init_document OK')
    return document

def set_primary_styles(document):
    """ Define genaral styles for all documents output
    TO DO: format YAML
    """
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.paragraph_format.space_after = Pt(0)
    
    styleArialTitle = document.styles.add_style('ArialTitle', WD_STYLE_TYPE.PARAGRAPH)
    styleArialTitle.font.name = 'Arial'
    styleArialTitle.font.size = Pt(24)
    styleArialTitle.font.bold = True
    styleArialTitle.paragraph_format.space_after = Pt(10)
    styleArialTitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    styleArialBlack = document.styles.add_style('ArialBlack', WD_STYLE_TYPE.CHARACTER)
    styleArialBlack.font.name = 'Arial Black'
    styleArialBlack.font.size = Pt(36)
    
    styleDesc = document.styles.add_style('Desc', WD_STYLE_TYPE.PARAGRAPH)
    styleDesc.font.name = 'Calibri'
    styleDesc.font.size = Pt(14)
    styleDesc.font.bold = True
    styleDesc.paragraph_format.space_after = Pt(10)
    styleDesc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    styleArialDesc = document.styles.add_style('ArialDesc', WD_STYLE_TYPE.CHARACTER)
    styleArialDesc.font.name = 'Arial'
    styleArialDesc.font.size = Pt(22)
    styleDesc.font.bold = False

    return document



def extract_tags(inpath, esp=''):
    """ Return a list of all tag elements contained in docx input
    """
    tree = get_tree(inpath)
    tagslist = []
    line = ''
    tab = 0
#    esp = '\t'
#    esp = '  '
    for elem in tostringlist(tree, encoding='unicode') :
        elem = elem.replace('≠','!=').replace('≤','=<').replace('≥','=>')
        elem = elem.replace('','???').replace('└','??')
        line += elem.replace('ns0:', '')
        if elem.startswith('</') :
            tab += -1
            tagslist.append(esp*tab + line)
            line = ''
        elif elem.endswith('/>') :
            tagslist.append(esp*tab + line)
            line = ''
        elif elem.endswith('>') :
            tagslist.append(esp*tab + line)
            tab += 1
            line = ''
        
    return tagslist


def write_json(dico, dicopath):
    """ Store data dict into json format
    """
    json_str = json.dumps(dico, sort_keys=True, indent=2)
    with open(dicopath, 'w', encoding="latin-1") as results:
        results.write(json_str)
                
    
def get_tree(path):
    """ Return docx content into xml format
    """
    document = zipfile.ZipFile(path)
    xml_content = document.read('word/document.xml')
    document.close()
    
    return XML(xml_content)

def set_cell_vertical_alignment(cell, align="center"): 
    """ Set cell's vertical alignment
    """
    try:   
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')  
        tcValign.set(qn('w:val'), align)  
        tcPr.append(tcValign)
        return True 
    except:
        traceback.print_exc()             
        return False
    
def set_cell_border(cell, **kwargs):
    """ Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
                    
                    
def add_tableOfContent(document):
    """ Add table of content in the document output
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    fldChar.set(qn('w:dirty'), 'true')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text  =  'TOC \\o \"1-6\" \\h \\z \\u'
    # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
#    p_element = paragraph._p
#    print(p_element)
    
    return document


def docxtoxml(inpath, outpath):
    """ Convert docx elements into a simpler readable list of elements with attributs
    Use : usefull to understand the structure of a new document before creating its parser
    """
    tree = extract_tags(inpath, esp='  ')
    with open(outpath, 'w') as results:
        for line in tree:
            # supprimer les caracteres non reconnus du docx dans le questionnaire CARE-S
            line = line.replace('↔', '<->').replace('', '->').replace('', ' ')
            results.write(line + '\n')
            
    logging.debug("docxtoxml OK")
    
    
    