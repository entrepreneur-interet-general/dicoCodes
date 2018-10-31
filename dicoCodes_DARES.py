# -*- coding: utf-8 -*-
"""
Created on Fri Jun  8 15:41:02 2018

@author: samah.ghalloussi
"""

import pandas as pd
import docutils
import json, regex, re, sys
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, RGBColor
from docx.shared import Pt, Cm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import logging
logging.basicConfig(level=logging.DEBUG)

startblue, endblue = "#startblue#", "/endblue/"

def extract_formatvar():
    """ Extract data from xlsx
        NOT USED anymore because data have been extracted and stored in another more convenient location
    """
    variables = pd.read_excel('Liste_Format_Variables.xlsx', 'Feuil1')
    cols = ['ID','Variable', 'Type', 'Len', 'Format', 'Informat', 'Label']
    variablesClean = pd.DataFrame(columns = variables.columns)

    variablesClean = variables[cols]
    first = variablesClean.to_dict(orient = 'index')
    second = {}
    for idx in first :
        try : second[first[idx]['Variable']] = first[idx]
        except : print('doublon :', idx)
    
def search_nearest(ffiltre, finfiltre, maxvalue=1):
    """ Return True if the nearest match for ffiltre is in finfiltre
        NOT USED anymore because of too much permissibility
    """
    text = '|'.join(finfiltre)
    #maximum permitted number of error: insertion + deletion + substitution
#        evalue = len(ffiltre.split(' ')) - 1
    evalue = 3
    for i in range(maxvalue) :
        #use the best match flag (?b) to get the single best match
        reg=regex.search('(?b)(?:'+ffiltre+'){e<='+str(evalue)+'}', text, re.IGNORECASE)
        if reg != None :
            return True
        else :
            evalue += 1
                
    return False

def comparprocess(text):
    """ Return normalized text to compare filters input with a level of permissibility
    """
    transf = text.lower().replace(' ou ','').replace(':','').replace('si','')
    
    if len(transf.split('(')[0]) > 5 :
        transf = transf.split('(')[0]
    else :
        transf = transf.replace('(','')
    
    return transf.replace(' ', '')
    
    
def actualise_filtre(filtre, finfiltre):
    """ Return the updated list of filters to apply to the next variable
    """
    # no need to actualize if no filter : return an empty list
    if filtre == [] :
        return filtre
    
    actufiltre = []
    nbf = 0
    for f in filtre :
        if comparprocess(f) not in finfiltre :#and not search_nearest(f.strip(), finfiltre) :
            actufiltre.append(f.strip())
        else :
            nbf += 1 
    
    #vérification de l'adéquation entre les filtres à supprimer et ceux supprimés
    if nbf != len(finfiltre) :
        logging.debug("Warning: number of Filters doesn't match") 

    return actufiltre


def check_filters(remainingfiltre, filtre, questvar) :
    """ Check filter doublons errors due to a FinFiltre written as a Filtre
    """
    doublons = []
    for f in remainingfiltre :
        if f in filtre :
            logging.debug("Problem: Filter already applied")
            logging.debug(f)
            logging.debug("Solution: Filter removed")
            doublons.append(comparprocess(f))
        
    return doublons

def parse_quest(tagslist):
    """ Return structured data from DARES questionnary input
    """
    current = {'tbl':0}
    i, r, c = 0, 0, 0
    table = {}
    tmp = ''
    tmpchoix = ''
    desc, paradesc = '', ''
    question = ''
    intable, soustable = False, False
    istitle = False
    firstlinetext = True
    isquestion = False
    isenum = False
    isgreen, isblue, isgray = False, False, False
    ignoreblue = False
    blocnb, varnb, enumnb = 0, 0, 0
    blocname = '--'
    sym=''
    nbfiltres = 0
    dicovar, dicobloc, dicoenum = {}, {}, {}
    enq, formatn, filtre, finfiltre = [], [], [], []
    isenq, isfiltre, isfinfiltre, isformat = False, False, False, False
    startfilter = False
    for line in tagslist:
        #replace symbols '≠' by '!=' :
        if line.startswith('<sym char=') :
            if "F0FF" in line :
                sym = ''
            elif "F0B3" in line :
                sym = '>='
            elif 'F0B9' in line :
                sym = '≠'
            elif "F0E0" in line :
                sym = ' --> '

        elif line.startswith('<p ') :
            firstlinetext = True
            isenum = False
        elif line == '</p>' :
            ignoreblue = False
            isgray = False
            # gere le cas où un finfiltre passe dans paradesc
            if paradesc.strip().startswith('Fin filtre') or paradesc.strip().startswith('Fin Filtre'):
                    finfiltre.append(paradesc.replace('Fin filtre','').replace('Fin Filtre', '').lstrip().lstrip(':').lstrip())
                    isfinfiltre = True
                    paradesc = ''
            if isfinfiltre :
                finfiltre[-1] = comparprocess(finfiltre[-1])
            isenq, isfiltre, isfinfiltre, isformat = False, False, False, False
            
            if intable :
                tmp += '\n'
            else :
                desc += '\n' + paradesc
                paradesc = ''
            desc = desc.replace('\n \n', '\n').replace('\n ', '\n').strip()
            
            if isquestion and '.' in question :
                #only for the first variable of a bloc
                if varnb == 0:
                    dicobloc[blocnb]['notes'] = desc.strip()
#                    filtre = []
                #all the other variables
                else :
                    nbfiltres += len(dicovar[varnb]['Filtre']) - len(finfiltre)
                    if nbfiltres > 0 :
                        remainingfiltre = actualise_filtre(dicovar[varnb]['Filtre'], finfiltre)
                        doublons = check_filters(remainingfiltre, filtre, dicovar[varnb]["Variable"])
                        #supprime les filtres en doublons
                        if doublons != [] :
                            remainingfiltre = actualise_filtre(remainingfiltre, doublons)
                            filtre = actualise_filtre(filtre, doublons)
                        filtre = remainingfiltre + filtre
                        nbfiltres -= len(remainingfiltre)  
                        
                    if desc.strip() != '' :
                        dicovar[varnb]['desc'] = desc
                        desc = ''
                    if enq != [] :
                        dicovar[varnb]['enq'] = enq
                    if formatn != [] :
                        dicovar[varnb]['format'] = formatn
                    if finfiltre != []:
                        dicovar[varnb]['FinFiltre'] = finfiltre
                    if dicoenum != {} :
                        dicovar[varnb]['choix'] = dicoenum
                dicoenum = {}
                enumnb = 0
                enq, formatn, finfiltre = [], [], []
                varnb += 1
                varname = question.replace('.', ' ').split()[0]
                question = question.replace(varname+'.', '').strip()
                dicovar[varnb] = {'Variable':varname, 'Question':question, 'choix': {}, 'Filtre':filtre, 'FinFiltre':finfiltre, 'desc':desc}
                desc = ''
                startfilter = False
                filtre = []
                question = ''
            
            else :
                paradesc += question + ' '
                question = ''
#                if startfilter :
#                    if varnb>0:
#                        dicovar[varnb]['Filtre'] = filtre
#                    startfilter = False

            istitle = False
            isquestion = False
                
        elif line == '<pStyle val="Titre1" />':
            istitle = True
#        elif line == '<pStyle val="Titre2" />':
#            istitle = True
#        elif line == '<pStyle val="Titre3" />':
#            istitle = True
#        elif line.startswith('<pStyle val="Titre'):
#            istitle = True
        elif line.startswith('<color val="008080" />') and firstlinetext :
            ignoreblue = True
        elif line.startswith('<color val="008080" />') and not ignoreblue and not firstlinetext :
            isblue = True
        elif line.startswith('<color val="00B050" />'):
            isgreen = True
        elif line.startswith('<color val="808080" />'):
            isgray = True
        elif line.startswith('<ilvl val='):
            isenum = True
            
        elif line == '<tbl>' :
            current['tbl'] += 1
            if not intable :
                intable = True
                r = 0
                table[i] = {}
            else :
                soustable = True

        elif line.startswith('<tr ') and not soustable :
            table[i][r] = {}
            c = 0
            
        elif line == '<tc>' and not soustable :
            tmp = ''

        elif line == '</tbl>' :
            current['tbl'] += -1
            if not soustable :
                intable = False
                i += 1
                table = {}
                r, c = 0, 0
            else :
                soustable = False
            
        elif line == '</tr>' and not soustable :
            r += 1
        
        elif line == '</tc>' and not soustable :
#            table[i]['table'][r][c] = tmp[0:-1]
            if c == 0 :
                tmpchoix = tmp.strip()
                if tmpchoix != '' :
                    #gere le tableau de la question QD32:
                    if not tmpchoix.startswith('Site '):
                        dicovar[varnb]['choix'][r+1] = [tmpchoix]
            elif tmp.strip() != '' and tmpchoix != '' and "……………………………" not in tmp :
                    dicovar[varnb]['choix'][r+1].append(tmp.strip())
#            if c == 0 :
#                tmpchoix = tmp.strip()
#                dicovar[varnb]['choix'][str(r)] = tmpchoix
#            elif c == 1 :
#                if tmp.strip() != '' and tmpchoix != '' and "……………………………" not in tmp :
#                    dicovar[varnb]['choix'][tmp.strip()] = tmpchoix
#                    del(dicovar[varnb]['choix'][str(r)])
#            elif c == 2 :
#                if tmp.strip() != '' and tmpchoix != '' and "……………………………" not in tmp :
#                    dicovar[varnb]['choix'][tmp.strip()] = tmpchoix
#                    del(dicovar[varnb]['choix'][str(r)])
            
#            elif tmp.strip() != '' :
#                table[i][r][c] = tmp.strip()
            tmp = ''
            c += 1
            
        elif line == '<tab />' :
            tmp += '\t'
            
        elif line.endswith('</t>') :
            text = line.replace('</t>', '').replace('\u00a0', ' ')
            if isblue :
                if not isfiltre and not isfinfiltre :
                    text = startblue + text + endblue
                isblue = False
                        
            if firstlinetext and not intable:
                if text.startswith('Q'+blocname) :
                    question = text
                    isquestion = True
                
            if istitle :
                if text.startswith('BLOC') :
                    if blocnb > 0 :
                        #ajout des infos de la derniere variable du precedant bloc
                        if varnb > 0 :
                            nbfiltres += len(dicovar[varnb]['Filtre']) - len(finfiltre)
                            if nbfiltres > 0 :
                                remainingfiltre = actualise_filtre(dicovar[varnb]['Filtre'], finfiltre)
                                doublons = check_filters(remainingfiltre, filtre, dicovar[varnb]["Variable"])
                                #supprime les filtres en doublons
                                if doublons != [] :
                                    remainingfiltre = actualise_filtre(remainingfiltre, doublons)
                                    filtre = actualise_filtre(filtre, doublons)
                                filtre = remainingfiltre + filtre
                                nbfiltres -= len(remainingfiltre)
                                
                            dicovar[varnb]['desc'] = desc.strip()
                            dicovar[varnb]['enq'] = enq
                            dicovar[varnb]['format'] = formatn
#                            filtre = actualise_filtre(filtre, finfiltre)
                            dicovar[varnb]['FinFiltre'] = finfiltre
                            enq, formatn, finfiltre = [], [], []
                        else :
                            dicobloc[blocnb]['notes'] = desc.strip()
                    #ajout des variables du precedant bloc
                    dicobloc[blocnb]['blocvar'] = dicovar
                    
                    #initialisation nouveau bloc
                    desc = ''
                    filtre = []
                    nbfiltres = 0
                    dicovar = {}
                    blocnb += 1
                    varnb = 0
                    try : blocname = text.split()[1]
                    except : blocname = '--'
                    dicobloc[blocnb] = {'bloctitle':text, 'blocname':blocname}
                
                elif blocnb > 0 :
                    dicobloc[blocnb]['bloctitle'] += text
                    
            elif isquestion and not firstlinetext :
#                dicovar[varnb]['question'] += text
                question += text
            elif intable :
                tmp += text
            
            elif isenum :
                if not isgray :
                    if firstlinetext :
                        enumnb += 1
                        dicoenum[enumnb] = [text]
                    else :
                        dicoenum[enumnb][-1] += text
                
            elif not isquestion :
                if text == 'VARIABLES DU QUESTIONNAIRE':# or text == 'Variables ajoutées par rapport aux variables collectées' or text == 'Variables de pondération':
                    blocname = ''
                    dicobloc[blocnb] = {'bloctitle':text, 'blocname':blocname}
                    filtre = []
                                
                elif text.startswith('ENQ :') or text.startswith('ENQ. :') or text.startswith('ENQ') :
                    enq.append(text.replace('ENQ :', '').replace('ENQ. :', '').lstrip())
                    isenq = True
                elif text.startswith('Format '):
                    formatn.append(text.replace('Format ',''))
                    isformat = True
                elif text.startswith('Fin filtre') or text.startswith('Fin Filtre') :#or text.startswith(startblue+'Fin filtre') or text.startswith(startblue+'Fin Filtre'):
#                    text = text.replace(startblue,'').replace(endblue,'')
                    finfiltre.append(text.replace('Fin filtre','').replace('Fin Filtre', '').lstrip().lstrip(':').lstrip())
#                    finfiltre.append(text)
                    isfinfiltre = True
                elif text.startswith('Filtre:') or text.startswith('Filtre :') or text.startswith('Filtre si '):
#                    text = text.replace(startblue,'').replace(endblue,'')
                    filtre.append(text.replace('Filtre','').replace(':', '').lstrip())
                    isfiltre = True
                    startfilter = True
                elif text == 'A tous' :
#                    filtre.append(text)
                    filtre = []
                    startfilter = True
                elif isenq :
                    enq[-1] = enq[-1] + sym + text
                    sym = ''
                elif isformat :
                    formatn[-1] = formatn[-1] + text
                elif isfinfiltre :
                    finfiltre[-1] = finfiltre[-1] + sym + text
                    sym = ''
                elif isfiltre:
                    filtre[-1] = filtre[-1] + sym + text
                    sym = ''                    
                else :
                    if not isgray :
                        paradesc += sym + text
                        sym = ''
                #stockage et réinitialisation des enumérations
                if dicoenum != {} :
                    if varnb > 0 :
                        dicovar[varnb]['choix'] = dicoenum
                    dicoenum = {}
                    enumnb = 0
                
            firstlinetext = False
    
    if blocnb > 0 :
        dicobloc[blocnb]['blocvar'] = dicovar
        dicobloc[blocnb]['notes'] = desc.strip()
      
    logging.debug("parse_quest OK")
    
    return dicobloc


def set_styles(document):
    """ Define specific styles for DARES document output
    TO DO: format YAML
    """
    styleDicoTitle = document.styles.add_style('dicoTitle', WD_STYLE_TYPE.PARAGRAPH)
    styleDicoTitle.font.name = 'Calibri'
    styleDicoTitle.font.size = Pt(16)
    styleDicoTitle.font.bold = True
    styleDicoTitle.font.underline = True
    styleDicoTitle.paragraph_format.space_after = Pt(0)
    styleDicoTitle.paragraph_format.widow_control = True
    
    styleBlocTitle = document.styles.add_style('blocTitle', WD_STYLE_TYPE.PARAGRAPH)
    styleBlocTitle.font.name = 'Calibri'
    styleBlocTitle.font.size = Pt(16)
    styleBlocTitle.paragraph_format.space_before = Pt(0)
    styleBlocTitle.paragraph_format.space_after = Pt(0)
    styleBlocTitle.paragraph_format.widow_control = True
    
    styleNormalCell = document.styles.add_style('normalCell', WD_STYLE_TYPE.PARAGRAPH)
    styleNormalCell.font.name = 'Calibri'
    styleNormalCell.font.size = Pt(10)
    styleNormalCell.paragraph_format.space_after = Pt(0)
    styleNormalCell.paragraph_format.widow_control = True
    styleNormalCell.paragraph_format.keep_with_next = True
    
    styleOrangeNormalCell = document.styles.add_style('orangeNormalCell', WD_STYLE_TYPE.PARAGRAPH)
    styleOrangeNormalCell.font.name = 'Calibri'
    styleOrangeNormalCell.font.size = Pt(10)
    styleOrangeNormalCell.font.color.rgb = RGBColor(228, 109, 10)
    styleOrangeNormalCell.paragraph_format.space_after = Pt(0)
    styleOrangeNormalCell.paragraph_format.widow_control = True
    styleOrangeNormalCell.paragraph_format.keep_with_next = True
    
    stylePurpleBoldCell = document.styles.add_style('purpleBoldCell', WD_STYLE_TYPE.PARAGRAPH)
    stylePurpleBoldCell.font.name = 'Calibri'
    stylePurpleBoldCell.font.size = Pt(10)
    stylePurpleBoldCell.font.color.rgb = RGBColor(112, 48, 160)
    stylePurpleBoldCell.paragraph_format.space_after = Pt(0)
    stylePurpleBoldCell.paragraph_format.keep_with_next = True
    stylePurpleBoldCell.paragraph_format.widow_control = True
    stylePurpleBoldCell.font.bold = True
    
    styleBlueNormalCell = document.styles.add_style('blueNormalCell', WD_STYLE_TYPE.PARAGRAPH)
    styleBlueNormalCell.font.name = 'Calibri'
    styleBlueNormalCell.font.size = Pt(10)
    styleBlueNormalCell.font.color.rgb = RGBColor(0, 112, 192)
    styleBlueNormalCell.paragraph_format.space_after = Pt(0)
    styleBlueNormalCell.paragraph_format.widow_control = True
    styleBlueNormalCell.paragraph_format.keep_with_next = True
           
    styleBoldNormalCell = document.styles.add_style('boldNormalCell', WD_STYLE_TYPE.PARAGRAPH)
    styleBoldNormalCell.font.name = 'Calibri'
    styleBoldNormalCell.font.size = Pt(10)
    styleBoldNormalCell.paragraph_format.space_before = Pt(0)
    styleBoldNormalCell.paragraph_format.space_after = Pt(0)
    styleBoldNormalCell.font.bold = True
    styleBoldNormalCell.paragraph_format.keep_with_next = True

    styleItalicNormalCell = document.styles.add_style('italicNormalCell', WD_STYLE_TYPE.PARAGRAPH)
    styleItalicNormalCell.font.name = 'Calibri'
    styleItalicNormalCell.font.size = Pt(10)
    styleItalicNormalCell.paragraph_format.space_before = Pt(3)
    styleItalicNormalCell.paragraph_format.space_after = Pt(0)
    styleItalicNormalCell.font.italic = True
    styleItalicNormalCell.paragraph_format.keep_with_next = True

    return document

def add_firstpage(document) :
    """ Add the firt page content in the document output
    """
    p = document.add_paragraph('\n\n', style='ArialTitle')
    p.add_run('\nDICTIONNAIRE DES VARIABLES\n\n', 'ArialBlack')
    document.add_paragraph('\n\nVersion du ' + docutils.get_date() + ' : générée automatiquement', style='Desc')

    document.add_page_break()
    
    document.add_paragraph('Dictionnaire des codes', style='dicoTitle')
    document.add_paragraph('\n', style='normalCell')
    
    return document



def add_filters(table, filters):
    """ Write filters information in table if needed
    """
    if filters != [] :
        for elemfilter in filters :
            cells = table.add_row().cells
            cells[0].merge(cells[1])
            docutils.cleanpara(cells[0]).add_paragraph(elemfilter, style='orangeNormalCell')
            fill_gray = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
            cells[0]._tc.get_or_add_tcPr().append(fill_gray)
    
    return table   


def add_dico(document, blocs) :
    """ Write data into the new document with DARES presentation
    """
    loggerBloc = logging.getLogger("Bloc")
    for numberbloc in blocs :
      if blocs[numberbloc]["blocname"] != '' :
        loggerBloc.debug(blocs[numberbloc]["blocname"])
        blocName = 'BLOC '+ blocs[numberbloc]["blocname"]
        document.add_paragraph(blocName, style='blocTitle')
        document.add_paragraph()

        for numbervar in blocs[numberbloc]["blocvar"] :
            table = document.add_table(rows=0, cols=2)
            # set column widths
            table.columns[0].width = Inches(2.0)
            table.columns[1].width = Inches(9.0)
            
            filters = blocs[numberbloc]["blocvar"][numbervar]["Filtre"]
            table = add_filters(table, filters)
                
            #ajout d'un ligne pour le nom de la variable et le libelle
            cells = table.add_row().cells
            #ajout du nom de la variable
            
            variable = blocs[numberbloc]["blocvar"][numbervar]["Variable"]
            docutils.cleanpara(cells[0]).add_paragraph(variable, style='purpleBoldCell')
            #ajout du libelle
#            docutils.cleanpara(cells[1]).add_paragraph(blocs[numberbloc]["blocvar"][numbervar]["Question"], style='boldNormalCell')
            rgb = RGBColor(0, 112, 192)
            docutils.add_color(rgb, startblue, endblue, docutils.cleanpara(cells[1]), blocs[numberbloc]["blocvar"][numbervar]["Question"], style='boldNormalCell')
            
            fill_gray = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
            cells[0]._tc.get_or_add_tcPr().append(fill_gray)
            fill_gray = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
            cells[1]._tc.get_or_add_tcPr().append(fill_gray)
            
            varquest = False
            choices = blocs[numberbloc]["blocvar"][numbervar]["choix"]
            if choices != {}:
                for elem in choices :
                    choix = choices[elem][0]
                    # pour eviter de réenumerer si la ligne est deja numerotée
                    if not choix.split('.')[0].strip().isdigit() or len(choix) < 4:
                        choix = str(elem) + ' - ' + choix
                    
                    if len(choices[elem]) == 1 :
                        cells = table.add_row().cells
                        fill_gray = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                        cells[0]._tc.get_or_add_tcPr().append(fill_gray)
                        fill_gray = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                        cells[1]._tc.get_or_add_tcPr().append(fill_gray)
#                        docutils.cleanpara(cells[1]).add_paragraph(choix, style='normalCell')
                        docutils.add_color(rgb, startblue, endblue, docutils.cleanpara(cells[1]), choix, style='normalCell')
                        

                    elif len(choices[elem]) == 2 :
                        #si variable par item, mettre la variable de la question en noir
                        if not varquest :
                            docutils.cleanpara(cells[0]).add_paragraph(variable, style='boldNormalCell')
                            varquest = True
                        cells = table.add_row().cells
                        fill_gray = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                        cells[0]._tc.get_or_add_tcPr().append(fill_gray)
                        fill_gray = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                        cells[1]._tc.get_or_add_tcPr().append(fill_gray)
#                        docutils.cleanpara(cells[1]).add_paragraph(choix, style='normalCell')
                        docutils.add_color(rgb, startblue, endblue, docutils.cleanpara(cells[1]), choix, style='normalCell')
                        docutils.cleanpara(cells[0]).add_paragraph(choices[elem][1], style='purpleBoldCell')
                    
                    else :
                        loggerBloc.debug(blocs[numberbloc]["blocvar"][numbervar]["Variable"])
                        loggerBloc.debug("Problem: more than 2 columns in tab : {} columns".format(len(choices[elem])) )
                        
            else :
                cells = table.add_row().cells
                fill_gray = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                cells[0]._tc.get_or_add_tcPr().append(fill_gray)
                fill_gray = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                cells[1]._tc.get_or_add_tcPr().append(fill_gray)
                desc = blocs[numberbloc]["blocvar"][numbervar]["desc"]
                docutils.cleanpara(cells[1]).add_paragraph(desc, style='normalCell')

            
            if cells[0].text != '' or cells[1].text != '' :
                cells = table.add_row().cells
                
            cells[0].merge(cells[1])
            docutils.cleanpara(cells[0]).add_paragraph(blocs[numberbloc]["blocvar"][numbervar]["FormatVar"], style='italicNormalCell')
            
            fill_gray = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
            cells[0]._tc.get_or_add_tcPr().append(fill_gray)
            fill_gray = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
            cells[1]._tc.get_or_add_tcPr().append(fill_gray)
            
#            ffilters = blocs[numberbloc]["blocvar"][numbervar]["FinFiltre"]
#            table = add_filters(table, ffilters)
            
#            cells = table.add_row().cells
            document.add_paragraph()
                
        document.add_page_break()
    
    return document

def set_margins(document, margin=2):
    """ Changing the page margins
    """
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    
    return document
    

def create_doc(blocs, outputdoc):
    """ Create a new document with DARES style
    """
    document = docutils.init_document()
    document = set_styles(document)
    document = add_firstpage(document)
    document = set_margins(document, margin=2)
    
    document = add_dico(document, blocs)
    
    document.save(outputdoc)
    logging.debug('create_doc OK')

def remplacements(text):
    """ Return text with proper symbols
    """
    text = text.replace('&lt;', '<').replace('&gt;', '>')
    text = text.replace('< ', '<').replace(' >', '>')
    text = text.replace('|_|_|','')
#    text = text.replace('|___________________|','')
    return text

def clean_data(blocs):
    """ Return cleaned dicodata
    """
    for numberbloc in blocs :
        blocs[numberbloc]["notes"] = remplacements(blocs[numberbloc]["notes"])
        for numbervar in blocs[numberbloc]["blocvar"] :
            blocs[numberbloc]["blocvar"][numbervar]["desc"] = remplacements(blocs[numberbloc]["blocvar"][numbervar]["desc"])
            blocs[numberbloc]["blocvar"][numbervar]["Question"] = remplacements(blocs[numberbloc]["blocvar"][numbervar]["Question"])
            
            cleanfilter = []
            for filtre in blocs[numberbloc]["blocvar"][numbervar]["Filtre"]:
                cleanfilter.append(remplacements(filtre))
            
            cleanfinfilter = []
            for finfiltre in blocs[numberbloc]["blocvar"][numbervar]["FinFiltre"]:
                cleanfinfilter.append(remplacements(finfiltre))
    
            blocs[numberbloc]["blocvar"][numbervar]["Filtre"] = cleanfilter
            blocs[numberbloc]["blocvar"][numbervar]["FinFiltre"] = cleanfinfilter
            
            for elem in blocs[numberbloc]["blocvar"][numbervar]["choix"]:
                blocs[numberbloc]["blocvar"][numbervar]["choix"][elem][0] = remplacements(blocs[numberbloc]["blocvar"][numbervar]["choix"][elem][0])
            
    return blocs


def merge_data(blocs, formatvar):
    """ Return final data with additionnal information from the formatvar input
    """
    for numberbloc in blocs :
        for numbervar in blocs[numberbloc]["blocvar"] :
            variable = blocs[numberbloc]["blocvar"][numbervar]["Variable"]
            choices = blocs[numberbloc]["blocvar"][numbervar]["choix"]
            fvar = ''
            #ajout d'un ligne Numérique si choix multiple sinon Caractère
            if variable == 'QB15' or variable == 'QB16' :
                fvar = "Une variable en 5 modalités pour chaque modalité de réponse (1 : Très important / 2 : Important / 3 : Peu important / 4 : Sans importance / 5 : Vous ne savez pas)"
            elif variable == 'QD42' or variable == 'QD62' :
                fvar = "Une variable en 3 modalités pour chaque modalité de réponse (1 : Oui / 2 : Non / 3 : Vous ne savez pas)"
            elif variable == 'QE5' or variable == 'QE9' :
                fvar = "Une variable en 3 modalités pour chaque modalité de réponse (1 : Oui, à tous / 2 : Oui, mais pas à tous  / 3 : Non, à aucun)"
            elif variable == 'QF16' :
                fvar = "Une variable en 3 modalités pour chaque modalité de réponse (1 : Avantage / 2 : Inconvénient / 3 : Ni l’un, ni l’autre)"
            elif choices != {} and len(choices[1]) != 1 :
                fvar = "Une variable dichotomique pour chaque modalité de réponse (0 : Oui / 1 : Non)"
#                docutils.cleanpara(cells[0]).add_paragraph("Une variable en 6 catégories pour chaque modalité de réponse (1 2 3 456)", style='italicNormalCell')
            else :
                try : fvar = formatvar[variable.upper()]
                except: 
                    try : fvar = formatvar['M'+variable.upper()+'M1']      
                    except : 
                        try: fvar = formatvar[variable.upper()+'_1']
                        except: 
                            try : fvar = formatvar[variable.upper()+'M1']
                            except : fvar = 'Caractère/Numérique'
                            
                if fvar == "Char" :
                    fvar = 'Caractère'
                elif fvar == "Num" :
                    fvar = 'Numérique'
                
                
                # ajout de l'info "Format" si presente
                try : tabformat = blocs[numberbloc]["blocvar"][numbervar]["format"]
                except : tabformat = []
                if tabformat != [] :
                    fvar += " : Format "+ ' '.join(tabformat)
            
            blocs[numberbloc]["blocvar"][numbervar]["FormatVar"] = fvar
            
    return blocs
    

def process_data(inputquest, inputvar):
    """ Return a dictionary with structured data from the questionnary and formatvar inputs
    """
    tagslist = docutils.extract_tags(inputquest)
    dicoquest = parse_quest(tagslist)
    dicobloc = clean_data(dicoquest) 
    with open(inputvar, 'r') as var_file:
        formatvar = json.load(var_file)
    dicodata = merge_data(dicobloc, formatvar)
    
    return dicodata

if __name__ == '__main__':
    # To run the program, type-in $ python dicoCodes_DARES.py [inputquest] [inputvar] [outpath]
    try :
        inputquest = sys.argv[1]
        inputvar = sys.argv[2]
        outpath = sys.argv[3]
        details = False #sys.argv[4]
    except :
        inputquest = "tests/inputs/Questionnaire_DARES.docx"
        inputvar = "tests/inputs/formatvariables_DARES.json"
        outpath = 'tests/outputs/'
        details = True

    if not outpath.endswith('/') :
        outpath += '/'
        
    dicoblocs = process_data(inputquest, inputvar)
    
    if details :
        docutils.docxtoxml(inputquest, outpath +inputquest.split('/')[-1]+"-xmlcontent.xml")
        docutils.write_json(dicoblocs, outpath + 'dicobloc-DARES.json')
        
    outputdoc = outpath + "autoDicoDARES-"+str(docutils.date.today())+".docx"
    create_doc(dicoblocs, outputdoc)