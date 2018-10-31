# -*- coding: utf-8 -*-
"""
Created on Tue Jun 12 13:20:38 2018

@author: samah.ghalloussi
"""

from __future__ import unicode_literals

import pandas as pd
import docutils
from openpyxl import load_workbook
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
import copy, os, sys
import win32com.client as win32
import os.path
import logging
logging.basicConfig(level=logging.DEBUG)

startred, endred = "#startred#", "/endred/"


           
def format_modalites(text, consignes) :
    """ Normalize modalites content
    """
    modalites = text.replace('vide =', 'vide\t=').replace('vide=', 'vide\t=')
    #uniformisation des espaces apres symboles
    modalites = modalites.replace('.', '. ').replace('=', '= ').replace('  ', ' ')
#    modalites = modalites.replace('=', '= ').replace('  ', ' ')
    #cas particulier à dissocier
    modalites = modalites.replace('. =', '===') 
    #remplacement des points par des tabulations
    modalites = modalites.replace('. ', ' \t')
    #cas particulier remis en forme
    modalites = modalites.replace('===', '. =')
    
    if "saisie en clair" in modalites :
        modalites = modalites.replace("saisie en clair","").strip()
        if "saisie en clair" not in consignes :
            consignes += "\nsaisie en clair"
            consignes = consignes.strip()
    
    return modalites.strip(), consignes
        
# separe le contenu du libellé de celui des consignes
def keep_consignes(content):
    """ Return values of elements from parsed text (separate consignes from libelle)
    """
    keeptab = ['(Plusieurs réponses possibles', '(plusieurs réponses possibles', '(une seule réponse possible)', '(libellé en clair', 'saisie en clair', 'Prendre en clair', 'INSERER TABLE PAYS']
    consignes = []   
    for keepconsigne in keeptab :
        if content.find(keepconsigne) != -1 :
            content = content.split(keepconsigne)[0]
            consignes.append(keepconsigne.replace('(', '').replace(')', ''))
    
    lines = content.split('\n')
    libelle = ''
    instructions = ''
    instructionfiltre = ''
    for line in lines :
        if line.strip().startswith("Instruction") or line.strip().startswith("[Citer liste") :
            instructions += line + '\n'
        elif line.strip().startswith("PROG") or line.strip().startswith("Exemples d’aide ") :
            instructions += line + '\n'
            
        elif line.strip().startswith("si N24=") : #handle N200 filter in libelle
            instructionfiltre += line + '\n'
        elif line.strip().lower().startswith("tendre carte") :
            instructions += line + '\n'
        elif line.strip().startswith("(plusieurs réponses possibles"):
            consignes.append(line.replace('(', '').replace(')', ''))
        else :
            libelle += '\n' + line
    
    libelle = libelle.replace('≠', '^=')
    questionmark = libelle.strip().split('?')
    
    instructions = ''
    
    if len(questionmark) == 1 :
        libelle = questionmark[0].strip()
    else :
        #gere plus de 1 question en libelle            
        libelle = '?'.join(questionmark[0:-1]).strip() + ' ?'
        instructions += questionmark[-1]
    
    
    if libelle.strip() == '' :
        libelle = instructionfiltre.strip()
        instructionfiltre = ''
    
    return '\n'.join(consignes), instructions.strip(), instructionfiltre.strip(), libelle


def parse_quest(tagslist):
    """ Return structured data from CARE-I questionnary input
    """
    dicoquest = {}
    i=0
    isborder, isintro, isinstruction, ismodule = False, False, False, False
    tempotext = ''
    intro, filtre = [], []
    finfiltre, varcalc = [], []
    question, variable, libelle = '', '', ''
    instructions, consignes, modalites = [], [], []
    other = []
    module = ''
    for line in tagslist:
        if line.startswith('<p ') :
            isborder = False
            isvarcalc = False
            isinstruction = False
            ismodule = False
        elif line.startswith('<color val="0000FF"') :
            isinstruction = True
        elif line.startswith('<pStyle val="TM1"') or line.startswith('<pStyle val="Titre2"'):
            ismodule = True
        elif line == '</p>' :
            tempotext = tempotext.strip()
            if ismodule :
                module = tempotext
                ismodule = False
            elif isborder :
                if tempotext.startswith('Filtre') :
                    filtre.append(tempotext)
                    isintro = False
                elif tempotext.startswith('Fin') :
                    finfiltre.append(tempotext)
                    isintro = False
                elif tempotext.startswith('Variable') or isvarcalc :
                    varcalc.append(tempotext)
                    isvarcalc = True
                elif ' - ' in tempotext:
                    module = tempotext

            elif tempotext.startswith('INTRO'):
                intro.append(tempotext)
                isintro = True
            elif ':' in tempotext :
                if ' - ' in tempotext and not tempotext.startswith('Si '):
                    if i>0:
                        dicoquest[i-1]["instructions"] = instructions
                        dicoquest[i-1]["other"] = other
                        dicoquest[i-1]["consignes"] = consignes
                        dicoquest[i-1]["finfiltre"] = finfiltre
                        dicoquest[i-1]["varcalc"] = varcalc
                        dicoquest[i-1]["modalites"] = modalites
                        instructions, consignes, modalites = [], [], []
                        finfiltre, varcalc = [], []
                        other = []
                        
                    question = tempotext.split()[0]
                    variable = tempotext.split()[2]
                    libelle = tempotext.split(':')[-1].strip()
                    dicoquest[i] = {
                                    "Module":module,
                                    "Question":question,
                                    "Variable":variable,
                                    "Libelle":libelle,
                                    "Filtre":filtre,
                                    "intro":intro
                                   }
                    i+=1
                    isintro = False
                    intro, filtre = [], []
                elif isintro :
                    intro.append(tempotext)
                elif i>0 and libelle == '' and tempotext.startswith('Si ') :
                    dicoquest[i-1]["Libelle"] += tempotext + '\n'
                else :
                    other.append(tempotext)
                    
            elif tempotext != '' :
                if i>0 and tempotext.startswith('\u2026') :
                    dicoquest[i-1]["Libelle"] += '\n' + tempotext
                elif i>0 and libelle == '' and tempotext.startswith('Si ') and not "modalité" in tempotext :
                    dicoquest[i-1]["Libelle"] += tempotext
                elif tempotext.startswith('(') or tempotext.startswith('NSP') :
                    consignes.append(tempotext)
                elif isinstruction and not tempotext.startswith('(') :
                    instructions.append(tempotext) 
                elif isintro :
                    intro.append(tempotext)
                else :
                    modalites.append(tempotext)
            
            tempotext = ''
                
        elif line == '<pBdr>' :
            isborder = True
            
        elif line.endswith('</t>') :
            text = line.replace('</t>', '').replace('\u00a0', ' ')
            tempotext += text
                
        
    return dicoquest

def process_quest(inpath, variables):
    """ Return a dictionary with structured data from the questionnary input
    """
    tagslist = docutils.extract_tags(inpath)
    dicoquest = parse_quest(tagslist, variables)
    
    return dicoquest

# ajoute les titres au document
def add_headings(document) :
    """ Add specific table of content titles in the document output
    """
    document.add_heading('I. Présentation de l’enquête', 1)
    document.add_heading('II. Échantillonnage et pondérations', 1)
    document.add_heading('III. Présentation des traitements aval', 1)
    document.add_page_break()
    document.add_section(WD_SECTION.CONTINUOUS)
    document.add_heading('IV. Dictionnaire des codes', 1)
    
    return document


def set_styles(document):
    """ Define specific styles for DARES document output
    TO DO: format YAML
    """
    heading2  = document.styles['Heading 2']
    heading2.paragraph_format.left_indent = Inches(0)
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    heading3  = document.styles['Heading 3']
    heading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    styleTitleTable = document.styles.add_style('titleTable', WD_STYLE_TYPE.PARAGRAPH)
    styleTitleTable.font.name = 'Calibri'
    styleTitleTable.font.size = Pt(20)
    styleTitleTable.font.bold = True
    styleTitleTable.paragraph_format.space_before = Pt(10)
    styleTitleTable.paragraph_format.space_after = Pt(10)
    styleTitleTable.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styleTitleTable.font.color.rgb = RGBColor(84, 141, 212)

    styleTitleModule = document.styles.add_style('titleModule', WD_STYLE_TYPE.PARAGRAPH)
    styleTitleModule.font.name = 'Calibri'
    styleTitleModule.font.size = Pt(16)
    styleTitleModule.font.bold = True
    styleTitleModule.font.color.rgb = RGBColor(84, 141, 212)
    styleTitleModule.paragraph_format.space_before = Pt(0)
    styleTitleModule.paragraph_format.space_after = Pt(10)
    styleTitleModule.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styleTitleModule.paragraph_format.widow_control = True
    
    styleBoldBiggerCell = document.styles.add_style('boldBiggerCell', WD_STYLE_TYPE.PARAGRAPH)
    styleBoldBiggerCell.font.name = 'Calibri'
    styleBoldBiggerCell.font.size = Pt(12)
    styleBoldBiggerCell.font.bold = True
    styleBoldBiggerCell.paragraph_format.keep_with_next = True
    styleBoldBiggerCell.paragraph_format.space_before = Pt(0)
    styleBoldBiggerCell.paragraph_format.space_after = Pt(0)
    
    styleBoldSmallCell = document.styles.add_style('boldSmallCell', WD_STYLE_TYPE.PARAGRAPH)
    styleBoldSmallCell.font.name = 'Calibri'
    styleBoldSmallCell.font.size = Pt(8)
    styleBoldSmallCell.font.bold = True
    styleBoldSmallCell.paragraph_format.keep_with_next = True
    styleBoldSmallCell.paragraph_format.space_before = Pt(0)
    styleBoldSmallCell.paragraph_format.space_after = Pt(0)
    
    styleBoldNormalCell = document.styles.add_style('boldNormalCell', WD_STYLE_TYPE.PARAGRAPH)
    styleBoldNormalCell.font.name = 'Calibri'
    styleBoldNormalCell.font.size = Pt(11)
    styleBoldNormalCell.font.bold = True
    styleBoldNormalCell.paragraph_format.keep_with_next = True
    styleBoldNormalCell.paragraph_format.space_before = Pt(3)
    styleBoldNormalCell.paragraph_format.space_after = Pt(0)
    
    styleBoldNormalCell = document.styles.add_style('italicNormalCell', WD_STYLE_TYPE.PARAGRAPH)
    styleBoldNormalCell.font.name = 'Calibri'
    styleBoldNormalCell.font.size = Pt(11)
    styleBoldNormalCell.font.italic = True
    styleBoldNormalCell.paragraph_format.keep_with_next = True
    styleBoldNormalCell.paragraph_format.space_before = Pt(0)
    styleBoldNormalCell.paragraph_format.space_after = Pt(0)
    
    styleConsignesCell = document.styles.add_style('ConsignesCell', WD_STYLE_TYPE.PARAGRAPH)
    styleConsignesCell.font.name = 'Calibri'
    styleConsignesCell.font.size = Pt(10)
    styleConsignesCell.font.italic = True
    styleConsignesCell.font.color.rgb = RGBColor(54, 95, 145)
    styleConsignesCell.paragraph_format.keep_with_next = True
    styleConsignesCell.paragraph_format.space_before = Pt(0)
    styleConsignesCell.paragraph_format.space_after = Pt(0)
    
    styleModalitesCell = document.styles.add_style('ModalitesCell', WD_STYLE_TYPE.PARAGRAPH)
    styleModalitesCell.font.name = 'Calibri'
    styleModalitesCell.font.size = Pt(10)
    styleModalitesCell.paragraph_format.keep_with_next = True
    styleModalitesCell.paragraph_format.space_before = Pt(3)
    styleModalitesCell.paragraph_format.space_after = Pt(3)
    styleModalitesCell.paragraph_format.widow_control = True
    
    styleFiltreCell = document.styles.add_style('FiltreCell', WD_STYLE_TYPE.PARAGRAPH)
    styleFiltreCell.font.name = 'Calibri'
    styleFiltreCell.font.size = Pt(10)
    styleFiltreCell.font.color.rgb = RGBColor(112, 48, 160)
    styleFiltreCell.paragraph_format.keep_with_next = True
    styleFiltreCell.paragraph_format.space_before = Pt(3)
    styleFiltreCell.paragraph_format.space_after = Pt(3)
    styleFiltreCell.paragraph_format.widow_control = True
    
    styleListVarTable = document.styles.add_style('listVarTable', WD_STYLE_TYPE.PARAGRAPH)
    styleListVarTable.font.name = 'Calibri'
    styleListVarTable.font.size = Pt(11)
    styleListVarTable.paragraph_format.keep_with_next = True
    styleListVarTable.paragraph_format.space_before = Pt(0)
    styleListVarTable.paragraph_format.space_after = Pt(10)
    
    styleListVarName = document.styles.add_style('listVarName', WD_STYLE_TYPE.PARAGRAPH)
    styleListVarName.font.name = 'Calibri'
    styleListVarName.font.size = Pt(9)
    styleListVarName.paragraph_format.keep_with_next = True
    styleListVarName.paragraph_format.space_before = Pt(0)
    styleListVarName.paragraph_format.space_after = Pt(0)
    
    return document

def add_firstpage(document) :
    """ Add the firt page content in the document output
    """
    p = document.add_paragraph('\n\n', style='ArialTitle')
    p.add_run('C', 'ArialBlack')
    p.add_run('apacités, ')
    p.add_run('A', 'ArialBlack')
    p.add_run('ides et ')
    p.add_run('RE', 'ArialBlack')
    p.add_run('ssources\ndes seniors\n')

    p1 = document.add_paragraph(style='ArialTitle')
    p1.add_run('\nCARE-Institutions') #.font.color.rgb = RGBColor(255, 0, 0)
#    p1.add_run('\n\n\nVolet « seniors »')
    p1.add_run('\n')
    p1.add_run('\n\nDictionnaire des codes\n\n')
    p2 = document.add_paragraph(style='Desc')
    p2.add_run('Version du ' + docutils.get_date() + ' : générée automatiquement').font.color.rgb = RGBColor(255, 0, 0)
    p3 = document.add_paragraph('\n\n', style='ArialTitle')
    p3.add_run('\nEnquête 2016', 'ArialDesc')
    document.add_page_break()
#    document.add_paragraph('Table des matières', style='ArialTitle')
    document = docutils.add_tableOfContent(document)
    document.add_page_break()
    
    
    return add_headings(document)

def ignore_var(dicovar, ignoretab):
    """ Remove the variables to ignore in the final document
    """
    dicovar_filtered = {}
    for idx in dicovar :
        if dicovar[idx]["Variable"] not in ignoretab :
            dicovar_filtered[idx] = dicovar[idx]
            
    return dicovar_filtered

# analyse le contenu d'un onglet du document xlsx et retourne un dico de variables
def xlsxParser(reader, table):
    dico = {}
    idx = 0
    module = ''
    filtre = []
    libelle = ''
    consignes = ''
    intro = ''
    
    for row in reader:
        row["v"] = ''
        row["Filtre"] = []
        row["Table"] = table
        if row['Libelle'] == '' :
            if row['Question'].lower().startswith('filtre') or row['Question'].lower().startswith('fin du filtre') or row['Question'].startswith('Remarque') or row['Question'].startswith('Début') :
                filtre.append(row['Question'])
            elif row['Question'].startswith('INTRO'):
                intro = row['Question'].split(':')[-1].strip()
            elif row['Variable'] == '' :
                filtre = []
                module = row['Question']
            else :
                idx += 1
                dico[idx] = row
                dico[idx]['Filtre'] = filtre
                dico[idx]['Module'] = module
                dico[idx]['intro'] = intro
                dico[idx]['Libelle'] = libelle + '\n- ' + dico[idx]['labelSAScourt']
                dico[idx]['Consignes'] = consignes
                dico[idx]["modalites"], consignes = format_modalites(dico[idx]["modalites"], consignes)
#                    longueur = ' (longueur = ' + str(row['longueur']) + ')'
                if row['type'] == "car" :
                    dico[idx]['type'] = 'Caractère'
                elif row['type'] == "num" :
                    dico[idx]['type'] = 'Numérique'
#                        [D1 Etab]  
        else :
            idx += 1
            dico[idx] = row
            dico[idx]['Module'] = module
            dico[idx]['intro'] = intro
#                longueur = ' (longueur = ' + str(row['longueur']) + ')'
            if row['type'] == "car" :
                dico[idx]['type'] = 'Caractère'
            elif row['type'] == "num" :
                dico[idx]['type'] = 'Numérique'
            
            consignes, instructions, instructionfiltre, libelle = keep_consignes(dico[idx]['Libelle'])
            dico[idx]["modalites"], consignes = format_modalites(dico[idx]["modalites"], consignes)
            
            if row['Question'] == 'N26' :
                libelle = libelle.split(' : ')[1]
            
            dico[idx]['instructions'] = instructions
            dico[idx]['Libelle'] = libelle
            dico[idx]['Consignes'] = consignes
            if instructionfiltre != "" :
                dico[idx]['Filtre'].append(instructionfiltre)
            
            dico[idx]['Variable'] = dico[idx]['Variable'].replace('- -','-').replace('--','-').replace(' ','')
            multivar = dico[idx]['Variable'].split('-')
            if len(multivar) == 1 :
                multivar = dico[idx]['Variable'].split('\n')
            
            labelSAS = dico[idx]['labelSAS'].strip().split('\n')
            multilibelle = dico[idx]['Libelle'].split('\n')
            if len(multivar) > 1 or dico[idx]['nbVariable'] == len(labelSAS) :
                if dico[idx]['nbVariable'] > 1 : #make sure only multiple variables here
                    basedico = copy.deepcopy(dico[idx])
                    dico[idx]['Variable'] = multivar[0]
                    for i, line in enumerate(labelSAS) :
                        values = line.split(' -')
                        if len(values) == 3 :
                            dico[idx] = copy.deepcopy(basedico)
                            dico[idx]['Variable'] = values[1].strip()
                            if i == 0 and dico[idx]['remarques'] != '' :
                                dico[idx]['remarques'] = startred + dico[idx]['remarques'].strip() + endred
                            else :
                                dico[idx]['remarques'] = ''
                            if len(multilibelle) == len(labelSAS) and dico[idx]['nbVariable'] > 1 : #####
                                dico[idx]['Libelle'] = multilibelle[i]
                            else :
                                dico[idx]['Libelle'] += '\n-' + values[2]
                            idx += 1                    
        
            if row['Question'] == 'N28' :
                dico[idx-1]['Libelle'] += '\n- ' + dico[idx-1]['labelSAScourt']
                    
    return dico

# regroupe les onglets d'un même dico ensemble et filtre les onglets à ignorer
def get_tablenames(sheet_names):
    
    tablenames = {}
    for name in sheet_names :
        if not name == 'ECHANTILLON':
            diconame = name.split('_')[0]
            #Renommer SENIORS_repondants en repondants
            if diconame == name :
                diconame = 'SENIORS'
            
            try :
                tablenames[diconame].append(name)
            except :
                tablenames[diconame] = [name]
    
    return tablenames


def add_variablesColumns(document, listAllVariables):
    """ Add the list of all the added variables at the end of the document
    in a 4 column way and return document
    """
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading('V. Liste des variables', 1)
    for tableName in listAllVariables :
        paratableName = document.add_paragraph("\nTable ", style='listVarTable')
        paratableName.add_run(tableName).font.bold = True
        document.add_section(WD_SECTION.CONTINUOUS)
        document.add_paragraph('\n'.join(listAllVariables[tableName]), style='listVarName')
        section = document.sections[-1]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'4')
        document.add_section(WD_SECTION.CONTINUOUS)
        section = document.sections[-1]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'1')
    
    return document


def list_modules(dicotables):
    """ Return a list of modules to use in footer
    """
    footermodules = []
    for tablename in dicotables:
        module = ''
        tablemodules = []
        dicovar = dicotables[tablename]
        for idx in dicovar :
            if dicovar[idx]["inclure"] == 1 :
                if module != dicovar[idx]["Module"] :
                    module = dicovar[idx]["Module"]
                    tablemodules.append('Table : ' + tablename + '\nModule : ' + module)
                    
        if tablemodules == [] :
            tablemodules.append('Table : '+ tablename)
        
        footermodules += tablemodules
            
    return footermodules

# ajoute le numéro et pied de page avec le nom de la table et du module
def add_footer(outputdoc, footermodules):
    """ Add pages number and footer with table and module names
    """
    cwd = os.getcwd()
    outfile = os.path.join(cwd, outputdoc)
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(outfile)
    word.Visible = False
    for i, module in enumerate(footermodules) :
        doc.Sections(i+1).Footers(1).LinkToPrevious = False
        doc.Sections(i+1).Footers(1).Range.Text = footermodules[i]
        doc.Sections(i+1).Footers(1).PageNumbers.Add()
    doc.Close(True)
    word.Application.Quit()


def add_dico(document, dicotables):
    """ Write CARE-I data from specific grouped tables into the new document 
    """
    loggerTable = logging.getLogger("Table")
    loggerIgnored = logging.getLogger("ignoredvar")
    for name in dicotables :
        loggerTable.debug(name)
        tablevar = dicotables[name]
        module = ''
        # cell border values by default
        cellbr = {"sz": 2, "val": "single", "color":"#808080"}
        document.add_heading('Table :\t'+ name, 2)
        for idx in tablevar :
            if tablevar[idx]["inclure"] == 1 :
                if module != tablevar[idx]["Module"] :
                    if module != '' :
                        document.add_page_break()
                        document.add_section(WD_SECTION.NEW_PAGE)
                        
                    module = tablevar[idx]["Module"]
                    document.add_heading('Module :\t' + module, 3)
                    
    
                    if tablevar[idx]["intro"] != '' :                    
                        intro = document.add_paragraph(style='boldNormalCell')
                        intro.add_run(tablevar[idx]["intro"]).font.bold = False
                        intro.paragraph_format.space_after = Pt(20)
                    else :
                        document.add_paragraph()
                
                else :
                    document.add_paragraph()
                        
                table = document.add_table(rows=0, cols=2)
                
                # set column widths
                table.columns[0].width = Inches(1.0)
                table.columns[1].width = Inches(7.0)
                
                cells = table.add_row().cells
                docutils.set_cell_border(cells[0], top=cellbr, bottom=cellbr, start=cellbr)
                docutils.set_cell_border(cells[1], top=cellbr, bottom=cellbr, end=cellbr)
    
                if len(tablevar[idx]["Question"])<10 and not tablevar[idx]["Question"].startswith("calcul"):
                    varquest = docutils.cleanpara(cells[0]).add_paragraph(style='boldBiggerCell')
                else :
                    varquest = docutils.cleanpara(cells[0]).add_paragraph(style='boldSmallCell')
                    docutils.set_cell_vertical_alignment(cells[0], align="center") 
        
                varquest.add_run(tablevar[idx]["Question"]).font.color.rgb = RGBColor(0, 0, 0)
                docutils.cleanpara(cells[1]).add_paragraph(tablevar[idx]["Variable"], style='boldBiggerCell')
                
                cells = table.add_row().cells
                cells[0].merge(cells[1])
                docutils.set_cell_border(cells[0], top=cellbr, start=cellbr, end=cellbr)
                if '\n-' in tablevar[idx]["Libelle"] :
                    #won't display italic 
                    libelle = tablevar[idx]["Libelle"].split('\n-')[0]
                    souslibelle = '\n-'.join(tablevar[idx]["Libelle"].split('\n-')[1::])
                    paralibelle = docutils.add_italic(docutils.cleanpara(cells[0]), libelle, style='boldNormalCell')
                    paralibelle.add_run('\n-' + souslibelle).font.bold = False
    
                else :
                    docutils.add_italic(docutils.cleanpara(cells[0]), tablevar[idx]["Libelle"], style='boldNormalCell')
        
                if tablevar[idx]["Consignes"].strip() != '' :
                    cells = table.add_row().cells
                    cells[0].merge(cells[1])
                    docutils.cleanpara(cells[0]).add_paragraph(tablevar[idx]["Consignes"], style='ConsignesCell')
                
                docutils.set_cell_border(cells[0], bottom=cellbr, start=cellbr, end=cellbr)            
                if tablevar[idx]["modalites"] != '' :
                    cells = table.add_row().cells
                    cells[0].merge(cells[1])
                    docutils.set_cell_border(cells[0], top=cellbr, bottom=cellbr, start=cellbr, end=cellbr)
                    docutils.cleanpara(cells[0]).add_paragraph(tablevar[idx]["modalites"], style='ModalitesCell')
                    
                if tablevar[idx]['type'] != 'Numérique' :
                    cells = table.add_row().cells
                    cells[0].merge(cells[1])
                    docutils.set_cell_border(cells[0], top=cellbr, bottom=cellbr, start=cellbr, end=cellbr)
                    docutils.cleanpara(cells[0]).add_paragraph(tablevar[idx]['type'], style='ModalitesCell')
        
                if tablevar[idx]["remarques"] != '' :
                    cells = table.add_row().cells
                    docutils.set_cell_border(cells[0], top=cellbr, bottom=cellbr, start=cellbr)
                    docutils.set_cell_border(cells[1], top=cellbr, bottom=cellbr, end=cellbr)
                    docutils.cleanpara(cells[0]).add_paragraph("Remarques", style='ModalitesCell')
                    rgb = RGBColor(54, 95, 145)
                    docutils.add_color(rgb, startred, endred, docutils.cleanpara(cells[1]), tablevar[idx]["remarques"], style='ModalitesCell')
                
                if tablevar[idx]["Filtre"] != [] :
                    cells = table.add_row().cells
                    docutils.set_cell_border(cells[0], top=cellbr, bottom=cellbr, start=cellbr)
                    docutils.set_cell_border(cells[1], top=cellbr, bottom=cellbr, end=cellbr)
                    docutils.cleanpara(cells[0]).add_paragraph("Filtre(s)", style='FiltreCell')
                    docutils.cleanpara(cells[1]).add_paragraph('\n'.join(tablevar[idx]["Filtre"]), style='FiltreCell')
                
            else :
                loggerIgnored.debug(tablevar[idx]["Variable"])
        
        document.add_section(WD_SECTION.CONTINUOUS)
        document.add_page_break()
        
    return document


def list_variables(dicotables) :
    """ Return a list of all the included variables of the document
    """
    dicovariables = {}
    for tablename in dicotables:
        tablevariables = []
        for idx in dicotables[tablename] :
            if dicotables[tablename][idx]["inclure"] == 1:
                tablevariables.append(dicotables[tablename][idx]["Variable"])
                
        dicovariables[tablename] = tablevariables
        
    return dicovariables

def add_variables(document, dicotables):
    """ Add the list of all the added variables at the end of the document
    in a table way and return the document updated
    """
    dicovariables = list_variables(dicotables)
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading('V. Liste des variables', 1)
    nbcols = 4
    for tablename in dicovariables :
        paratableName = document.add_paragraph("\nTable ", style='listVarTable')
        paratableName.add_run(tablename).font.bold = True
        length = len(dicovariables[tablename])
        nbrows = int(length / nbcols)
        if length % nbcols != 0 :
            nbrows += 1
        table = document.add_table(rows=nbrows, cols=nbcols)
        idx = 0
        for j in range(nbcols) :
            for i in range(nbrows) :
                if idx < length :
                    docutils.cleanpara(table.cell(i, j)).add_paragraph(dicovariables[tablename][idx], style='listVarName')
                    idx += 1
    
    return document

def create_doc(dicotables, outputdoc):
    """ Create a new document with CARE-I style
    """
    document = docutils.init_document()
    document = set_styles(document)
    document = add_firstpage(document)
    
    document = add_dico(document, dicotables)
    
    document = add_variables(document, dicotables)
    
    footermodules = list_modules(dicotables)
    footermodules += ['Liste des variables']
    
    document.save(outputdoc)
    
    add_footer(outputdoc, footermodules)
    
    logging.debug('create_doc OK')
    

def parse_xlsx(inpath):
    """ Return a dictionary with extracted data from input xlsx document sheets
    """
    workbook = load_workbook(inputxlsx, data_only=True)
    sheet_names = workbook.sheetnames
    xlsxreader = pd.read_excel(open(inpath,'rb'), header=0, sheet_name=sheet_names)
#    check_validxlsx(xlsxreader)
    alldico = {}
    for sheetname in xlsxreader :
        reader = {}
        for label in xlsxreader[sheetname]:
            for i, cell in enumerate(xlsxreader[sheetname][label]):
                if pd.isnull(cell) :
                    cell = ''
                if label in ['inclure','nbVariable'] and cell != '':
                    cell = int(cell)
                try : 
                    reader[i][label] = cell
                except : 
                    reader[i] = {}
                    reader[i][label] = cell
        rdico = []
        for row in reader:
            if reader[row]['Question'] != '' :
                rdico.append(reader[row])
        
        alldico[sheetname] = xlsxParser(rdico, sheetname)
    
    return alldico


def process_xlsx(dicoxlsx):
    """ Return a dictionary with structured data from the xlsx document input
    """
    dicodata = {}
    for tablename in dicoxlsx :
        if not tablename == 'ECHANTILLON':
            diconame = tablename.split('_')[0]
            #Renommer SENIORS_repondants en repondants
            if diconame == tablename :
                diconame = 'SENIORS'
            
            
            try :
                dicodata[diconame][tablename] = dicoxlsx[tablename]
            except :
                dicodata[diconame] = {}
                dicodata[diconame][tablename] = dicoxlsx[tablename]
                 
    return dicodata
    

def process_data(inputxlsx, inputquest):
    """ Return a dictionary with structured data from the specs and questionnary inputs
    """
    dicoxlsx = parse_xlsx(inputxlsx)
    dicodata = process_xlsx(dicoxlsx)
    
    dicoquests = {}
    for diconame in dicodata :
        inpath = inputquest.replace('[]',diconame)
        tagslist = docutils.extract_tags(inpath)
        dicoquests[diconame] = parse_quest(tagslist)
        
        
    return dicodata


if __name__ == '__main__':
    # To run the program, type-in $ python dicoCodes_CAREI.py [inputspec] [inputquest] [outpath]
    try :
        inputxlsx = sys.argv[1]
        inputquest = sys.argv[2]
        outpath = sys.argv[3]
        details = False #sys.argv[4]
    except :
        inputxlsx = 'CARE-I/ALL/CARE-I Specs_pourDicoCodes_ALL_181003.xlsx'
        inputquest = 'CARE-I/CARE-I Questionnaires (word)/Questionnaire '+'[]'+'_171127_pourSiteDrees.docx'
        outpath = 'tests/outputs/'
        details = True

    if not outpath.endswith('/') :
        outpath += '/'
        
    dicodata = process_data(inputxlsx, inputquest)

    if details :
        docutils.write_json(dicodata, outpath + 'dicodata-CAREI.json') 
        for diconame in dicodata :
            inputdoc = inputquest.replace('[]', diconame)
            output = outpath + inputdoc.split('/')[-1] + "-xmlcontent.xml"
            docutils.docxtoxml(inputdoc, output)
    
    mydate = str(docutils.date.today()).replace('-', '')[2::]
    for diconame in dicodata :
        outputdoc = outpath + 'autoDicoCARE-I_'+ diconame + '_'+ mydate + '.docx' 
        create_doc(dicodata[diconame], outputdoc)
        

